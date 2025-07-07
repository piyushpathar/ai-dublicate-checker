from flask import Flask, render_template, request, Response, stream_with_context
from werkzeug.utils import secure_filename
from docx import Document
import fitz
import re
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import json
import time

app = Flask(__name__)
model = SentenceTransformer('sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2')

def clean_gujarati_text(text):
    return re.sub(r"[\"',।!?\\\-]", "", text).strip()

def is_valid_line(line):
    if not line.strip():
        return False
    if re.match(r"^(Page|પાનું)\s*\d+\s*(of|/)\s*\d+", line, re.IGNORECASE):
        return False
    if re.sub(r"\s+", "", line).isalnum() is False and re.sub(r"\s+", "", line) == "":
        return False
    return True

def split_gujarati_questions(text):
    split_lines = re.split(r"(પ્રશ્ન\s*\d+[\.\:\)]|Que\.\s*\d+[\.\:\)]?)", text)
    cleaned, buffer = [], ""
    for part in split_lines:
        if re.match(r"(પ્રશ્ન\s*\d+[\.\:\)]|Que\.\s*\d+[\.\:\)]?)", part):
            if buffer.strip():
                cleaned.append(clean_gujarati_text(buffer))
            buffer = part
        else:
            buffer += " " + part
    if buffer.strip():
        cleaned.append(clean_gujarati_text(buffer))
    return [line for line in cleaned if len(line.split()) >= 3 and is_valid_line(line)]

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_file():
    file = request.files.get('file')
    if not file or file.filename == '':
        return "No file selected", 400
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[-1].lower()

    def generate():
        extracted = []
        if ext == 'docx':
            doc = Document(file)
            total = len(doc.paragraphs)
            for idx, p in enumerate(doc.paragraphs):
                line = p.text.strip()
                if line and is_valid_line(line):
                    for question in split_gujarati_questions(line):
                        extracted.append((question, "Page ~"))
                yield f"data: {json.dumps({'progress': f'Reading paragraph {idx+1}/{total}'})}\n\n"
                time.sleep(0.01)
        elif ext == 'pdf':
            doc_pdf = fitz.open(stream=file.read(), filetype="pdf")
            total = doc_pdf.page_count
            for i, page in enumerate(doc_pdf, start=1):
                text = page.get_text()
                for line in text.split('\n'):
                    line = line.strip()
                    if line and is_valid_line(line):
                        for question in split_gujarati_questions(line):
                            extracted.append((question, f"Page {i}"))
                yield f"data: {json.dumps({'progress': f'Reading page {i}/{total}'})}\n\n"
                time.sleep(0.01)
        else:
            yield f"data: {json.dumps({'error': 'Unsupported file type'})}\n\n"
            return

        # Now embedding + similarity
        sentences, pages = zip(*extracted) if extracted else ([], [])
        if len(sentences) > 0:
            yield f"data: {json.dumps({'progress': f'Encoding {len(sentences)} sentences'})}\n\n"
            embeddings = model.encode(list(sentences), batch_size=64)
            sim_matrix = cosine_similarity(embeddings)
            duplicates, seen = [], set()
            for i in range(len(sentences)):
                if i in seen:
                    continue
                for j in range(i + 1, len(sentences)):
                    if j in seen:
                        continue
                    if sim_matrix[i][j] >= 0.88:
                        duplicates.append((
                            sentences[i], pages[i],
                            sentences[j], pages[j],
                            round(sim_matrix[i][j], 4)
                        ))
                        seen.add(j)
            yield f"data: {json.dumps({'done': True, 'duplicates': duplicates})}\n\n"
        else:
            yield f"data: {json.dumps({'done': True, 'duplicates': []})}\n\n"

    return Response(stream_with_context(generate()), mimetype='text/event-stream')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=3000, threaded=True)
